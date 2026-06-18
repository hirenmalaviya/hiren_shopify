# -*- coding: utf-8 -*-
"""Generate the client-facing step-by-step flow document as a .docx."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "screenshots")
OUT = os.path.join(BASE, "CSV-to-Shopify-Import-Guide.docx")

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x05, 0x96, 0x69)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK

for lvl, sz in ((1, 20), (2, 15), (3, 12)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = INDIGO if lvl < 3 else DARK
    st.font.bold = True

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.9)
sec.top_margin = sec.bottom_margin = Inches(0.9)
CONTENT_W = Inches(6.6)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text="", *, size=11, bold=False, italic=False, color=None, align=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def screenshot(name, caption):
    path = os.path.join(SHOTS, name)
    if not os.path.exists(path):
        para(f"[missing screenshot: {name}]", color=RED, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=CONTENT_W)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def callout(text):
    """A shaded note box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, "EEF2FF")
    cell.width = CONTENT_W
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def kvtable(rows, headers):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(val).font.size = Pt(10)
    return tbl


# ======================= TITLE =======================
para("CSV → Shopify Product Import System", size=26, bold=True, color=INDIGO,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=40)
para("Step-by-Step User Guide", size=16, bold=True, color=DARK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("A complete walkthrough of how to use the application — with real screenshots",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para("Laravel 12  ·  Vue 3 + Inertia  ·  PrimeVue  ·  Shopify Admin GraphQL API",
     size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=INDIGO, space_after=2)
para("Every screenshot in this document is a real capture of the running application "
     "(captured 18 June 2026 against the live Shopify sandbox).",
     size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

doc.add_page_break()

# ======================= OVERVIEW =======================
doc.add_heading("1. What this application does", level=1)
para("You upload a Shopify-style product CSV. The application validates the file (in the browser "
     "and again on the server), stores it, and creates a database record for every row. A background "
     "queue then imports each product into a Shopify store through the Admin GraphQL API — creating "
     "the product, its variant (price, SKU, weight, inventory policy), setting inventory at a "
     "location, attaching the image, and adding it to the target collection.")
para("Throughout the import, a live dashboard shows each product moving through "
     "pending → processing → successful / failed, with full logging, an in-app log viewer, and "
     "failure notifications. Re-uploading the same file updates the existing products instead of "
     "creating duplicates.")

doc.add_heading("How the system is put together", level=2)
para("The flow from a click in the browser to a product in Shopify:")
flow = (
    "Browser (Vue 3 + Inertia + PrimeVue)\n"
    "      │  upload CSV\n"
    "      ▼\n"
    "UploadController  →  validate type, size, required headers\n"
    "      ▼\n"
    "CsvImportService  →  parse + per-row validation + duplicate detection\n"
    "                     → 1 Upload record + N Product records\n"
    "      ▼\n"
    "Database queue  →  ProcessCsvImport  →  one job per product\n"
    "      ▼\n"
    "ShopifyService (GraphQL)\n"
    "      create | update  →  variant  →  inventory  →  image  →  collection\n"
    "      ▼\n"
    "status + Shopify IDs + logs + notifications  →  live dashboard"
)
codep = doc.add_paragraph()
cr = codep.add_run(flow)
cr.font.name = "Consolas"
cr.font.size = Pt(9)
codep.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ======================= STEPS =======================
doc.add_heading("2. Step-by-step walkthrough", level=1)

# Step 1
doc.add_heading("Step 1 — The upload page", level=2)
para("Open the application at http://localhost:8000. You land on the import page.")
screenshot("01-upload-page.png", "The import page — drag-and-drop upload with a format guide.")
bullets([
    ("A drag-and-drop dropzone", " (or the Browse files button)."),
    ("A format guide", " on the right listing the required columns (Handle, Title, Variant SKU, "
     "Variant Price) and the optional ones."),
    ("A Download CSV template link", " so you always start from a correctly-formatted file."),
    ("Top navigation", " — Upload · Dashboard · Logs, plus a notification bell and a dark-mode toggle."),
])

# Step 2
doc.add_heading("Step 2 — Select a CSV", level=2)
para("Drop in (or browse to) your product file. The app shows the file name and size, and the "
     "Import products button becomes active.")
screenshot("02-file-selected.png", "A file is selected — name and size are shown, the button is enabled.")
callout("The sample file used here contains 10 valid products.")

# Step 3
doc.add_heading("Step 3 — Client-side validation", level=2)
para("If you pick the wrong kind of file (or one that is too big or empty), the app rejects it "
     "instantly in the browser, before anything is uploaded.")
screenshot("11-client-validation.png", "Client-side validation rejects a non-CSV file immediately.")
para("Client-side validation checks the file type (.csv only) and size (≤ 5 MB) and rejects empty "
     "files. The server then re-validates everything independently and additionally confirms that "
     "the required column headers exist — so the rules can never be bypassed.")

# Step 4
doc.add_heading("Step 4 — Upload & live processing", level=2)
para("Click Import products. You get a success message — “Upload received — 10 rows queued for "
     "import” — and land on the detail page, which immediately starts tracking progress.")
screenshot("03-detail-processing.png", "The detail page tracks each product live as it is imported.")
bullets([
    ("A progress bar", " (0 / 10 processed) and live counters: successful / failed / skipped."),
    ("A per-product table", " — row #, Title, SKU, Price, Status, Action, Detail."),
    ("Auto-refresh", " is on — the page updates every few seconds without a manual refresh."),
])

# Step 5
doc.add_heading("Step 5 — Completion", level=2)
para("Within a few seconds every product reaches a final state and the upload is marked Completed.")
screenshot("04-detail-completed.png", "All 10 products imported successfully — 100% complete.")
bullets([
    ("10 / 10 processed · 100%", " — 10 successful, 0 failed, 0 skipped."),
    ("The Action column shows “update” here", " because these products already existed in the "
     "Shopify sandbox from earlier runs — the system matched and updated them instead of creating "
     "duplicates (the update-existing / upsert feature). On a brand-new store these would read "
     "“create”."),
    ("View in Shopify links", " open each product directly in the Shopify admin."),
])

# Step 6
doc.add_heading("Step 6 — The dashboard", level=2)
para("The Dashboard gives the high-level overview of every import.")
screenshot("05-dashboard.png", "The dashboard — stat cards, a status chart, and the uploads table.")
bullets([
    ("Stat cards", " — Total uploads, Products imported, Failed, Success rate."),
    ("A “Products by status” donut chart", "."),
    ("An All uploads table", " with search, sortable columns, pagination, a per-upload progress bar, "
     "and view / delete actions."),
])

# Step 7
doc.add_heading("Step 7 — The log viewer", level=2)
para("Every import event is recorded. The Logs page is an in-app viewer over those records — you can "
     "filter by level and search messages.")
screenshot("06-logs.png", "The log viewer — every import event, filterable and searchable.")
para("Each row expands to reveal the full JSON context captured for that event:")
screenshot("07-logs-expanded.png",
           "An expanded log row showing the Shopify product ID, SKU, handle, and row number.")

# Step 8
doc.add_heading("Step 8 — Validation & failure handling", level=2)
para("Now upload a deliberately bad CSV — a row with a non-numeric price, a row with a missing "
     "handle, and a row with a missing title.")
screenshot("08-bad-csv-failed.png", "Each invalid row is marked Failed with a precise, readable error.")
bullets([
    ("The upload is marked Failed", "; the counters read 0 successful, 3 failed."),
    ("Each row carries a precise error", " — e.g. “Variant Price must be a number ≥ 0 (got "
     "“not-a-number”).”, “Handle is required.”, “Title is required.”"),
    ("Validation is per-row", " — one bad row never aborts the whole file."),
    ("The notification bell", " now shows a red unread badge, and Retry failed / Delete actions appear."),
])

# Step 9
doc.add_heading("Step 9 — Failure notifications", level=2)
para("Click the bell. An in-app, database-backed notification summarises the failure.")
screenshot("09-notification.png", "An in-app notification summarises the failed import.")
para("“Import failed — 3 of 3 products failed to import in ‘bad-products.csv’.” "
     "with a timestamp and a Mark all read action.")

# Step 10
doc.add_heading("Step 10 — Dark mode", level=2)
para("The whole application has a polished dark theme, toggled from the top-right and remembered "
     "between visits. Here is the dashboard after both imports — the success rate has dropped to 77% "
     "and the chart now shows the failed slice in red.")
screenshot("10-dashboard-dark.png", "The dashboard in dark mode after a successful and a failed import.")

doc.add_page_break()

# ======================= REQUIREMENTS MAP =======================
doc.add_heading("3. How each requirement is demonstrated", level=1)
kvtable([
    ("CSV upload form, clean responsive UI", "Steps 1–2"),
    ("Client-side validation (type & size)", "Step 3"),
    ("Server-side validation + required headers", "Steps 3–4, 8"),
    ("Asynchronous (queued) processing", "Steps 4–5"),
    ("Per-product status (pending→processing→successful/failed)", "Steps 4–5"),
    ("Shopify integration via GraphQL (create/variant/inventory/image)", "Step 5"),
    ("Add to required collection 464337174767", "Step 5"),
    ("Update existing product (upsert)", "Step 5 (update action)"),
    ("Dashboard: all imports, stats, chart, search/sort/paginate", "Step 6"),
    ("Error messages for failed rows", "Step 8"),
    ("Logging + in-app log viewer", "Step 7"),
    ("Error notification system", "Steps 8–9"),
    ("Retry & delete", "Steps 6, 8"),
    ("Dark mode / world-class UI polish", "Step 10 (and all)"),
], ["Requirement", "Shown in"])

# ======================= REPRODUCE =======================
doc.add_heading("4. How to run it yourself", level=1)
para("Two processes are needed — the web server and the queue worker. Run each command in its own "
     "terminal from the project folder.")
steps = [
    "# 1. Install dependencies and build the front-end",
    "composer install",
    "npm install",
    "npm run build",
    "",
    "# 2. Configure .env (database + Shopify keys), then migrate",
    "php artisan migrate",
    "",
    "# 3. Verify the Shopify connection (token, location, collection)",
    "php artisan shopify:check",
    "",
    "# 4. Start the app (two terminals)",
    "php artisan serve        # → http://localhost:8000",
    "php artisan queue:work   # processes the imports",
]
cp = doc.add_paragraph()
cr = cp.add_run("\n".join(steps))
cr.font.name = "Consolas"
cr.font.size = Pt(9.5)
cp.paragraph_format.space_after = Pt(10)

para("Sample files:", bold=True, space_after=2)
bullets([
    ("Valid CSV (10 products): ", "plan/shopify-products-csv (5) (1) (1).csv"),
    ("Invalid CSV (for the failure demo): ", "docs/sample-csv/bad-products.csv"),
])

para("For a full requirement-by-requirement mapping to the source files, see "
     "docs/CLIENT_REQUIREMENTS.md. For a narrated video version of this flow, see docs/DEMO_VIDEO.md.",
     italic=True, color=GREY, space_before=8)

doc.save(OUT)
print("WROTE:", OUT)
